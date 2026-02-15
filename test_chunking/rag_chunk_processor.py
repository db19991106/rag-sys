#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
RAG知识库文档处理器 - 混合切分版本
功能：数据清洗 → 结构切分 → 语义切分 → 嵌入计算 → 结构化JSON输出

处理流程：
1. 数据清洗：读取多格式文档，去除无效字符，规整表格，统一编码
2. 结构切分：按章节/子标题/条款拆分，保留表格结构
3. 语义切分：使用BGE模型计算相似度，相似度≥0.8合并，<0.8拆分
4. 结构化输出：生成指定格式的JSON文件
"""

import json
import os
import re
import math
import hashlib
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from dataclasses import dataclass, field, asdict
import warnings
warnings.filterwarnings('ignore')

import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity


# ==================== 数据类定义 ====================

@dataclass
class TableData:
    """表格数据结构"""
    table_id: str
    rows: List[List[str]]
    chapter: str = ""
    subtitle: str = ""
    
    def to_dict(self) -> Dict:
        return {
            "table_id": self.table_id,
            "rows": self.rows,
            "chapter": self.chapter,
            "subtitle": self.subtitle
        }


@dataclass
class StructuredUnit:
    """结构切分后的基本单元"""
    chapter: str = ""  # 章节名称
    subtitle: str = ""  # 子标题
    articles: List[str] = field(default_factory=list)  # 条款列表
    content: str = ""  # 纯文本内容
    tables: List[TableData] = field(default_factory=list)  # 表格数据
    char_count: int = 0  # 字符数


@dataclass
class Chunk:
    """最终切分后的chunk"""
    chunk_id: str
    chapter: str
    subtitle: str
    articles: List[str]
    content: str
    tables: List[Dict]
    metadata: Dict[str, Any]
    embeddings: Optional[np.ndarray] = None  # 临时存储，不序列化
    
    def to_dict(self) -> Dict:
        return {
            "chunk_id": self.chunk_id,
            "chapter": self.chapter,
            "subtitle": self.subtitle,
            "articles": self.articles,
            "content": self.content,
            "tables": self.tables,
            "metadata": self.metadata
        }


# ==================== 1. 数据清洗模块 ====================

class DataCleaner:
    """数据清洗器：处理各种格式文档，去除无效字符，规整表格"""
    
    def __init__(self):
        self.encoding_errors = []
        
    def clean_text(self, text: str) -> str:
        """
        清洗文本内容
        - 去除乱码和无效字符
        - 规整空白符
        - 统一编码
        """
        if not text:
            return ""
        
        # 1. 统一转换为UTF-8（已读取为str，无需重复解码）
        # 2. 去除控制字符（保留换行和制表符）
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text)
        
        # 3. 去除零宽字符
        text = re.sub(r'[\u200b-\u200f\ufeff]', '', text)
        
        # 4. 规整空白符：多个空格/制表符转为单个空格
        text = re.sub(r'[ \t]+', ' ', text)
        
        # 5. 规整换行：多个连续换行保留最多两个
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        
        # 6. 去除每行首尾空白
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # 7. 去除页眉页脚类内容
        text = re.sub(r'\d+\s*/\s*\d+\s*页', '', text)
        text = re.sub(r'第\s*\d+\s*页', '', text)
        text = re.sub(r'Page\s*\d+\s*(of|/)\s*\d+', '', text, flags=re.IGNORECASE)
        
        # 8. 去除常见的页眉内容
        page_header_patterns = [
            r'WORD格式可编辑',
            r'专业知识整理分享',
            r'Word文档',
            r'可编辑',
        ]
        for pattern in page_header_patterns:
            text = re.sub(pattern, '', text)
        
        # 9. 去除无关内容（文档末尾的垃圾内容）
        garbage_patterns = [
            r'单纯的课本内容，并不能满足',
            r'儿童画',
            r'狐狸和鸡',
            r'小鸭子学游泳',
            r'后悔也来不及',
            r'摘草莓的小姑娘',
            r'学生的需要，通过补充',
            r'儿童意愿画',
            r'儿童扩散性思维',
            r'情节，都是一则有趣的小故事',
        ]
        for pattern in garbage_patterns:
            text = re.sub(pattern, '', text)
        
        # 10. 去除单独存在的短句（可能是页眉页脚残留）
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            # 跳过太短的行（除非是有效的编号）
            if len(line) < 5 and not re.match(r'^[（(]?\d+[）)]?\s*$', line):
                continue
            cleaned_lines.append(line)
        text = '\n'.join(cleaned_lines)
        
        # 10. 去除URL和水印
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', '', text)
        
        return text.strip()
    
    def extract_tables_from_text(self, text: str) -> Tuple[str, List[TableData]]:
        """
        从文本中提取表格
        支持Markdown表格、ASCII表格、制表符分隔等格式
        """
        tables = []
        table_id_counter = 1
        
        # 模式1: Markdown表格
        md_pattern = r'\|(.+)\|\n\|[-:\s|]+\|\n((?:\|.+\|\n?)+)'
        md_matches = list(re.finditer(md_pattern, text))
        
        for match in md_matches:
            header_line = match.group(1)
            body_lines = match.group(2).strip().split('\n')
            
            # 解析表头
            headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
            
            # 解析行
            rows = [headers]
            for line in body_lines:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    rows.append(cells)
            
            table = TableData(
                table_id=f"T1_{table_id_counter:03d}",
                rows=rows
            )
            tables.append(table)
            table_id_counter += 1
            
            # 从文本中移除表格
            text = text.replace(match.group(0), f"\n[TABLE_{table.table_id}]\n")
        
        # 模式2: 简单ASCII表格（空格/制表符对齐）
        # 检测连续多行具有相似结构的文本
        lines = text.split('\n')
        table_ranges = []
        in_table = False
        table_start = 0
        
        for i, line in enumerate(lines):
            # 检测表格行：包含多个空格分隔的内容
            if re.match(r'^[\s\u4e00-\u9fa5a-zA-Z0-9]+(\s{2,}[\s\u4e00-\u9fa5a-zA-Z0-9]+){2,}$', line):
                if not in_table:
                    in_table = True
                    table_start = i
            else:
                if in_table and i - table_start >= 2:  # 至少2行才认为是表格
                    table_ranges.append((table_start, i))
                in_table = False
        
        # 处理检测到的表格
        for start, end in reversed(table_ranges):  # 反向处理以便删除
            table_lines = lines[start:end]
            if len(table_lines) >= 2:
                # 简单分割
                rows = []
                for line in table_lines:
                    cells = [cell.strip() for cell in re.split(r'\s{2,}', line) if cell.strip()]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    table = TableData(
                        table_id=f"T1_{table_id_counter:03d}",
                        rows=rows
                    )
                    tables.append(table)
                    table_id_counter += 1
                    
                    # 替换为占位符
                    lines[start:end] = [f"[TABLE_{table.table_id}]"]
        
        text = '\n'.join(lines)
        
        return text, tables
    
    def load_document(self, filepath: Path) -> Tuple[str, List[TableData], List[int]]:
        """
        加载文档，支持TXT、DOC、DOCX、PDF格式
        返回：(清洗后的文本, 表格列表, 表格段落索引列表)
        """
        suffix = filepath.suffix.lower()
        
        if suffix == '.txt':
            text, tables = self._load_txt(filepath)
            return text, tables, []
        elif suffix == '.doc':
            return self._load_doc(filepath)
        elif suffix == '.docx':
            return self._load_docx(filepath)
        elif suffix == '.pdf':
            text, tables = self._load_pdf(filepath)
            return text, tables, []
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
    
    def _load_txt(self, filepath: Path) -> Tuple[str, List[TableData]]:
        """加载TXT文件"""
        # 尝试不同编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin1']
        raw_text = None
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    raw_text = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if raw_text is None:
            raise ValueError(f"无法解码文件: {filepath}")
        
        # 清洗文本
        cleaned_text = self.clean_text(raw_text)
        
        # 提取表格
        text, tables = self.extract_tables_from_text(cleaned_text)
        
        return text, tables
    
    def _load_doc(self, filepath: Path) -> Tuple[str, List[TableData], List[int]]:
        """加载DOC文件（旧版Word格式或伪装成doc的docx）"""
        # 方法1：尝试用python-docx直接读取（很多.doc实际是docx格式）
        try:
            from docx import Document
            doc = Document(filepath)
            
            # 创建段落和表格对象的ID映射
            para_by_elem = {p._element: p for p in doc.paragraphs}
            table_by_elem = {t._element: t for t in doc.tables}
            
            result_elements = []  # (type, content)
            tables = []
            table_counter = 1
            
            # 遍历body的所有子元素，根据tag判断类型
            for child in doc.element.body:
                tag = child.tag.lower()
                
                if tag.endswith('p'):  # 段落
                    para = para_by_elem.get(child)
                    if para:
                        text = para.text.strip()
                        if text:
                            result_elements.append(('para', text))
                elif tag.endswith('tbl'):  # 表格
                    table = table_by_elem.get(child)
                    if table:
                        rows = []
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            rows.append(cells)
                        
                        if rows:
                            table_id = f"T1_{table_counter:03d}"
                            table_data = TableData(
                                table_id=table_id,
                                rows=rows
                            )
                            tables.append(table_data)
                            result_elements.append(('table', table_id))
                            table_counter += 1
            
            # 构建文本：段落和表格交替
            text_parts = []
            for elem_type, content in result_elements:
                if elem_type == 'para':
                    text_parts.append(content)
                elif elem_type == 'table':
                    text_parts.append(f"[TABLE_{content}]")
            
            text = '\n'.join(text_parts)
            cleaned_text = self.clean_text(text)
            
            # 返回表格索引
            table_indices = []
            para_count = 0
            for elem_type, _ in result_elements:
                if elem_type == 'para':
                    para_count += 1
                elif elem_type == 'table':
                    table_indices.append(para_count)
            
            return cleaned_text, tables, table_indices
        except Exception as e:
            pass
        
        # 方法2：尝试使用antiword命令
        try:
            import subprocess
            result = subprocess.run(['antiword', str(filepath)], 
                                  capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                raw_text = result.stdout
                cleaned_text = self.clean_text(raw_text)
                text, tables = self.extract_tables_from_text(cleaned_text)
                return text, tables, []
        except (subprocess.TimeoutExpired, FileNotFoundError, Exception):
            pass
        
        # 方法3：尝试使用textract库
        try:
            import textract
            raw_text = textract.process(str(filepath), encoding='utf-8').decode('utf-8')
            cleaned_text = self.clean_text(raw_text)
            text, tables = self.extract_tables_from_text(cleaned_text)
            return text, tables, []
        except Exception as e:
            print(f"  textract处理失败: {e}")
        
        # 方法4：尝试使用docx2txt
        try:
            import docx2txt
            raw_text = docx2txt.process(str(filepath))
            cleaned_text = self.clean_text(raw_text)
            text, tables = self.extract_tables_from_text(cleaned_text)
            return text, tables, []
        except Exception as e:
            print(f"  docx2txt处理失败: {e}")
        
        # 方法5：使用olefile提取文本（纯文本方式）
        try:
            import olefile
            if olefile.isOleFile(str(filepath)):
                ole = olefile.OleFileIO(str(filepath))
                if ole.exists('WordDocument'):
                    word_stream = ole.openstream('WordDocument').read()
                    text_parts = []
                    i = 0
                    while i < len(word_stream):
                        if 32 <= word_stream[i] <= 126 or word_stream[i] >= 128:
                            try:
                                char = bytes([word_stream[i]]).decode('utf-8', errors='ignore')
                                if char:
                                    text_parts.append(char)
                            except:
                                pass
                        elif word_stream[i] in (0x0d, 0x0a):
                            text_parts.append('\n')
                        i += 1
                    raw_text = ''.join(text_parts)
                    cleaned_text = self.clean_text(raw_text)
                    text, tables = self.extract_tables_from_text(cleaned_text)
                    return text, tables, []
        except Exception as e:
            print(f"  olefile处理失败: {e}")
        
        # 所有方法都失败
        raise ValueError(
            f"无法处理DOC文件: {filepath.name}\n"
            f"请安装以下任一依赖：\n"
            f"  - antiword (系统包): apt-get install antiword\n"
            f"  - textract: pip install textract\n"
            f"  - docx2txt: pip install docx2txt\n"
            f"  - olefile: pip install olefile"
        )
    
    def _load_docx(self, filepath: Path) -> Tuple[str, List[TableData], List[int]]:
        """
        加载DOCX文件，返回(文本, 表格列表, 表格段落索引列表)
        """
        try:
            from docx import Document
        except ImportError:
            print("警告: python-docx未安装，跳过DOCX处理")
            return "", [], []
        
        doc = Document(filepath)
        
        # 创建段落和表格对象的ID映射
        para_by_elem = {p._element: p for p in doc.paragraphs}
        table_by_elem = {t._element: t for t in doc.tables}
        
        result_elements = []  # (type, content) - type: 'para' or 'table'
        tables = []
        table_counter = 1
        
        # 遍历body的所有子元素，根据tag判断类型
        for child in doc.element.body:
            tag = child.tag.lower()
            
            if tag.endswith('p'):  # 段落
                para = para_by_elem.get(child)
                if para:
                    text = para.text.strip()
                    if text:
                        result_elements.append(('para', text))
            elif tag.endswith('tbl'):  # 表格
                table = table_by_elem.get(child)
                if table:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)
                    
                    if rows:
                        table_id = f"T1_{table_counter:03d}"
                        table_data = TableData(
                            table_id=table_id,
                            rows=rows
                        )
                        tables.append(table_data)
                        result_elements.append(('table', table_id))
                        table_counter += 1
        
        # 构建文本：段落和表格交替
        text_parts = []
        for elem_type, content in result_elements:
            if elem_type == 'para':
                text_parts.append(content)
            elif elem_type == 'table':
                text_parts.append(f"[TABLE_{content}]")
        
        text = '\n'.join(text_parts)
        cleaned_text = self.clean_text(text)
        
        # 返回表格索引（每个表格在第几个段落之后）
        table_indices = []
        para_count = 0
        for elem_type, _ in result_elements:
            if elem_type == 'para':
                para_count += 1
            elif elem_type == 'table':
                table_indices.append(para_count)
        
        return cleaned_text, tables, table_indices
    
    def _load_pdf(self, filepath: Path) -> Tuple[str, List[TableData]]:
        """加载PDF文件"""
        try:
            import PyPDF2
        except ImportError:
            print("警告: PyPDF2未安装，跳过PDF处理")
            return "", []
        
        text_parts = []
        tables = []
        
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        raw_text = '\n'.join(text_parts)
        cleaned_text = self.clean_text(raw_text)
        
        # 尝试提取表格（简化处理）
        text, tables = self.extract_tables_from_text(cleaned_text)
        
        return text, tables


# ==================== 2. 结构切分模块 ====================

class StructureParser:
    """结构解析器：按章节/条款进行精准切分"""
    
    # 章节标题正则 - 只匹配「第X章」格式
    CHAPTER_PATTERN = r'^(第[一二三四五六七八九十百千]+章|第\d+章)\s*'
    CHAPTER_REGEX = re.compile(CHAPTER_PATTERN)
    
    # 条款正则 - 只匹配「第一条」「第1条」格式
    ARTICLE_PATTERN = r'^(第[一二三四五六七八九十百千]+条|第\d+条)\s*'
    ARTICLE_REGEX = re.compile(ARTICLE_PATTERN)
    
    def __init__(self):
        pass
    
    def is_chapter_start(self, line: str) -> bool:
        """检查是否是章节开始"""
        return bool(self.CHAPTER_REGEX.match(line.strip()))
    
    def is_article_start(self, line: str) -> bool:
        """检查是否是条款开始"""
        return bool(self.ARTICLE_REGEX.match(line.strip()))
    
    def identify_line_type(self, line: str) -> str:
        """识别行类型：章节标题、子标题或正文内容"""
        line = line.strip()
        if not line:
            return "empty"
        if self.is_chapter_start(line):
            return "chapter"
        if self.is_article_start(line):
            return "subtitle"
        return "content"
    
    def extract_subtitle_number(self, line: str) -> str:
        """提取子标题编号（如"第一条"、"（一）"等）"""
        line = line.strip()
        if self.ARTICLE_REGEX.match(line):
            return line
        return ""
    
    def extract_chapter(self, line: str) -> str:
        """提取章节完整标题（如"第一章 总则"）"""
        match = self.CHAPTER_REGEX.match(line.strip())
        if match:
            return match.group(0).strip()
        return ""
    
    def extract_article_number(self, line: str) -> str:
        """提取条款编号（如"第一条"）"""
        match = self.ARTICLE_REGEX.match(line.strip())
        if match:
            return match.group(0).strip()
        return ""
    
    def parse_structure(self, text: str, tables: List[TableData], table_indices: List[int] = None) -> List[StructuredUnit]:
        """
        解析文档结构，生成结构化单元列表
        策略：
        1. 按「第一层：按「第一章」切分章节
        2. 在章节内按「第一层：按「第一条」切分条款
        3. 每个条款作为一个基础单元
        """
        if table_indices is None:
            table_indices = []
        
        lines = text.split('\n')
        units = []
        
        # 创建表格ID到表格对象的映射
        table_map = {t.table_id: t for t in tables}
        
        # 状态变量
        current_chapter = ""
        current_article = ""  # 当前条款号
        content_buffer = []  # 当前条款内容缓存
        current_tables = []  # 当前条款关联的表格
        
        def flush_article():
            """保存当前条款为一个unit"""
            nonlocal units, current_chapter, current_article, content_buffer, current_tables
            
            if not content_buffer:
                return
            
            content = '\n'.join(content_buffer).strip()
            if content:
                # 关联表格内容
                table_text = ""
                for t in current_tables:
                    table_text += f"\n\n【表格 {t.table_id}】\n"
                    for row in t.rows:
                        table_text += " | ".join(row) + "\n"
                
                full_content = content + table_text
                
                unit = StructuredUnit(
                    chapter=current_chapter if current_chapter else "未分类",
                    subtitle=current_article,
                    articles=[current_article] if current_article else [],
                    content=full_content,
                    char_count=len(full_content),
                    tables=[t for t in current_tables]
                )
                units.append(unit)
            
            content_buffer = []
            current_tables = []
        
        # 跟踪当前处理到的段落索引
        current_para_idx = 0
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # 非空行计数
            if line:
                current_para_idx += 1
            
            if self.is_chapter_start(line):
                # 新章节开始，保存之前的条款
                flush_article()
                current_chapter = line.strip()
                current_article = ""
                
            elif self.is_article_start(line):
                # 新条款开始，保存之前的条款
                flush_article()
                current_article = line.strip()
                content_buffer = [line]
                
            elif line:
                # 检查是否是表格标记
                table_match = re.match(r'\[TABLE_(T\d+_\d+)\]', line)
                if table_match:
                    table_id = table_match.group(1)
                    if table_id in table_map:
                        table = table_map[table_id]
                        table.chapter = current_chapter
                        current_tables.append(table)
                else:
                    # 累积到当前条款内容
                    content_buffer.append(line)
            
            i += 1
        
        # 保存最后一个条款
        flush_article()
        
        return units
    
    def parse_structure(self, text: str, tables: List[TableData], table_indices: List[int] = None) -> List[StructuredUnit]:
        """
        解析文档结构，生成结构化单元列表
        核心策略：以子标题/条款为最小单元，每个独立一个unit
        表格根据位置关联到当前章节
        """
        if table_indices is None:
            # 从清洗后的文本中重新计算表格位置
            lines = text.split('\n')
            table_indices = []
            for i, line in enumerate(lines):
                if '[TABLE_' in line:
                    table_indices.append(i)
        
        lines = text.split('\n')
        units = []
        
        # 创建表格ID到表格对象的映射
        table_map = {t.table_id: t for t in tables}
        
        # 状态变量
        current_chapter_title = ""  # 章节完整标题（含"第X章"）
        current_subtitle = ""  # 当前子标题编号（如"（一）"）
        content_buffer = []  # 当前unit内容缓存
        current_tables = []  # 当前unit关联的表格
        
        # 跟踪当前处理到的行索引
        current_line_idx = 0
        
        def flush_content():
            """保存当前内容为一个unit"""
            nonlocal units, current_chapter_title, current_subtitle, content_buffer, current_tables
            
            if not content_buffer and not current_tables:
                return
            
            content = '\n'.join(content_buffer).strip() if content_buffer else ""
            
            # 如果有关联表格，将表格内容添加到content中
            if current_tables:
                table_content = "\n\n【表格】\n"
                for t in current_tables:
                    table_content += f"\n{t.table_id}:\n"
                    for row in t.rows:
                        table_content += " | ".join(row) + "\n"
                content = content + table_content if content else table_content
            
            if content:
                unit = StructuredUnit(
                    chapter=current_chapter_title if current_chapter_title else "未分类",
                    subtitle=current_subtitle,
                    articles=[current_subtitle] if current_subtitle else [],
                    content=content,
                    char_count=len(content),
                    tables=[t for t in current_tables]
                )
                units.append(unit)
            
            content_buffer = []
            current_tables = []
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            line_type = self.identify_line_type(line)
            
            # 更新当前行索引（非空行）
            if line:
                current_line_idx += 1
            
            if line_type == "chapter":
                # 先保存之前的内容
                flush_content()
                
                # 开始新章节
                current_chapter_title = line.strip()
                
            elif line_type == "subtitle":
                # 遇到新的子标题，先保存之前的
                flush_content()
                
                # 开始新的子标题
                current_subtitle = self.extract_subtitle_number(line)
                content_buffer = [line]
                
            elif line_type == "content":
                if line:
                    # 检查是否是表格标记
                    table_match = re.match(r'\[TABLE_(T\d+_\d+)\]', line)
                    if table_match:
                        # 先保存当前内容
                        flush_content()
                        
                        # 创建单独的表格unit，关联到当前章节
                        table_id = table_match.group(1)
                        if table_id in table_map:
                            table = table_map[table_id]
                            table.chapter = current_chapter_title
                            # 创建一个只包含表格的unit
                            table_unit = StructuredUnit(
                                chapter=current_chapter_title,
                                subtitle="",
                                articles=[],
                                content=f"[表格]",
                                char_count=0,
                                tables=[table]
                            )
                            units.append(table_unit)
                    else:
                        # 累积到当前内容
                        content_buffer.append(line)
            
            i += 1
        
        # 保存最后一个内容
        flush_content()
        
        return units
        
        # 保存最后一个内容
        flush_content()
        
        return units
    
    def _merge_short_units(self, units: List[StructuredUnit], min_chars: int = 100) -> List[StructuredUnit]:
        """合并过短的连续单元"""
        if not units:
            return units
        
        merged = []
        current = units[0]
        
        for i in range(1, len(units)):
            next_unit = units[i]
            
            # 如果当前单元太短且属于同一章节/子标题，尝试合并
            if (current.char_count < min_chars and 
                current.chapter == next_unit.chapter and
                current.subtitle == next_unit.subtitle and
                not current.tables and not next_unit.tables):  # 不包含表格
                
                # 合并内容
                current.content += '\n' + next_unit.content
                current.articles.extend(next_unit.articles)
                current.char_count = len(current.content)
            else:
                merged.append(current)
                current = next_unit
        
        merged.append(current)
        return merged


# ==================== 3. 语义切分模块 ====================

class SemanticChunker:
    """语义切分器：使用BGE模型计算语义相似度"""
    
    def __init__(self, model_path: str, similarity_threshold: float = 0.8):
        """
        初始化语义切分器
        
        Args:
            model_path: BGE模型路径
            similarity_threshold: 相似度阈值，≥此值则合并
        """
        self.model_path = model_path
        self.similarity_threshold = similarity_threshold
        self.model = None
        self._load_model()
    
    def _load_model(self):
        """加载BGE嵌入模型"""
        try:
            print(f"正在加载嵌入模型: {self.model_path}")
            self.model = SentenceTransformer(self.model_path)
            print(f"✅ 模型加载成功")
        except Exception as e:
            print(f"❌ 模型加载失败: {e}")
            print("将使用备用方案（不进行语义切分）")
            self.model = None
    
    def compute_embedding(self, text: str) -> np.ndarray:
        """计算文本的嵌入向量"""
        if self.model is None:
            # 返回零向量作为后备
            return np.zeros(768)
        
        # 限制文本长度，避免过长
        if len(text) > 1000:
            text = text[:1000]
        
        embedding = self.model.encode(text, convert_to_numpy=True)
        return embedding
    
    def compute_similarity(self, emb1: np.ndarray, emb2: np.ndarray) -> float:
        """计算两个嵌入向量的余弦相似度"""
        if self.model is None:
            return 1.0  # 后备方案：假设高相似度
        
        # 归一化
        emb1_norm = emb1 / (np.linalg.norm(emb1) + 1e-8)
        emb2_norm = emb2 / (np.linalg.norm(emb2) + 1e-8)
        
        # 计算余弦相似度
        similarity = np.dot(emb1_norm, emb2_norm)
        return float(similarity)
    
    def semantic_chunking(self, units: List[StructuredUnit], chapter_num: int) -> List[Chunk]:
        """
        对结构切分后的单元进行语义切分
        
        策略：
        1. 计算每个单元的嵌入向量
        2. 遍历单元，计算相邻单元间的相似度
        3. 相似度≥阈值则合并，<阈值则拆分
        4. 每个chunk记录平均语义相似度
        """
        if not units:
            return []
        
        # 计算所有单元的嵌入
        print(f"  计算 {len(units)} 个单元的语义嵌入...")
        embeddings = []
        for i, unit in enumerate(units):
            emb = self.compute_embedding(unit.content)
            embeddings.append(emb)
        
        # 语义聚类与合并
        chunks = []
        current_chunk_units = [units[0]]
        current_chunk_embeddings = [embeddings[0]]
        similarities = []
        chunk_idx = 1
        
        for i in range(1, len(units)):
            current_unit = units[i]
            current_emb = embeddings[i]
            
            # 计算与当前chunk中所有单元的平均相似度
            if current_chunk_embeddings:
                sims = [self.compute_similarity(current_emb, emb) for emb in current_chunk_embeddings]
                avg_sim = sum(sims) / len(sims)
                similarities.append(avg_sim)
                
                # 决策：是否合并
                if avg_sim >= self.similarity_threshold:
                    # 合并到当前chunk
                    current_chunk_units.append(current_unit)
                    current_chunk_embeddings.append(current_emb)
                else:
                    # 保存当前chunk，开始新chunk
                    chunk = self._create_chunk(
                        current_chunk_units, 
                        current_chunk_embeddings,
                        chapter_num,
                        chunk_idx,
                        similarities if similarities else [1.0]
                    )
                    chunks.append(chunk)
                    
                    # 重置
                    current_chunk_units = [current_unit]
                    current_chunk_embeddings = [current_emb]
                    similarities = []
                    chunk_idx += 1
        
        # 处理最后一个chunk
        if current_chunk_units:
            chunk = self._create_chunk(
                current_chunk_units,
                current_chunk_embeddings,
                chapter_num,
                chunk_idx,
                similarities if similarities else [1.0]
            )
            chunks.append(chunk)
        
        return chunks
    
    def _create_chunk(
        self, 
        units: List[StructuredUnit], 
        embeddings: List[np.ndarray],
        chapter_num: int,
        chunk_idx: int,
        similarities: List[float]
    ) -> Chunk:
        """从单元列表创建Chunk"""
        
        # 合并内容
        contents = []
        articles = []
        tables = []
        chapter = units[0].chapter if units else ""
        subtitle = units[0].subtitle if units else ""
        
        for unit in units:
            if unit.content:
                contents.append(unit.content)
            articles.extend(unit.articles)
            tables.extend(unit.tables)
        
        merged_content = '\n\n'.join(contents)
        
        # 计算平均相似度
        avg_similarity = sum(similarities) / len(similarities) if similarities else 1.0
        avg_similarity = round(avg_similarity, 3)
        
        # 确定切分原因
        if avg_similarity >= self.similarity_threshold:
            chunking_reason = f"强语义关联(相似度{avg_similarity:.3f})"
        else:
            chunking_reason = f"弱语义关联(相似度{avg_similarity:.3f})"
        
        # 生成chunk_id
        chunk_id = f"CH{chapter_num}_{chunk_idx:03d}"
        
        # 统计信息
        char_count = len(merged_content)
        article_count = len(articles)
        has_table = len(tables) > 0
        
        # 构建metadata
        metadata = {
            "char_count": char_count,
            "article_count": article_count,
            "avg_similarity": avg_similarity,
            "has_table": has_table,
            "chunking_reason": chunking_reason
        }
        
        # 转换tables为dict列表
        tables_dict = [t.to_dict() for t in tables]
        
        return Chunk(
            chunk_id=chunk_id,
            chapter=chapter,
            subtitle=subtitle,
            articles=articles,
            content=merged_content,
            tables=tables_dict,
            metadata=metadata
        )


# ==================== 4. 主处理器 ====================

class RAGDocumentProcessor:
    """RAG文档处理器：整合所有模块，执行完整处理流程"""
    
    def __init__(
        self, 
        data_dir: str,
        output_dir: str,
        model_path: str = "backend/data/models/BAAI--bge-base-zh-v1.5",
        similarity_threshold: float = 0.8
    ):
        """
        初始化处理器
        
        Args:
            data_dir: 输入数据目录
            output_dir: 输出目录
            model_path: BGE模型路径
            similarity_threshold: 语义相似度阈值
        """
        self.data_dir = Path(data_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 初始化各模块
        self.cleaner = DataCleaner()
        self.parser = StructureParser()
        self.chunker = SemanticChunker(model_path, similarity_threshold)
        
        # 全局chunk_id计数器
        self.chunk_id_counter = 0
        self.used_chunk_ids = set()
    
    def extract_chapter_number(self, chapter_name: str) -> int:
        """从章节名称提取数字"""
        # 尝试匹配"第X章"
        match = re.search(r'第([一二三四五六七八九十百千\d]+)章', chapter_name)
        if match:
            num_str = match.group(1)
            # 转换中文数字
            chinese_nums = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
                          '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
            if num_str in chinese_nums:
                return chinese_nums[num_str]
            elif num_str.isdigit():
                return int(num_str)
        
        # 尝试匹配"X、"
        match = re.search(r'^([一二三四五六七八九十])[、．.]', chapter_name)
        if match:
            chinese_nums = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5,
                          '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
            return chinese_nums.get(match.group(1), 1)
        
        return 1
    
    def generate_unique_chunk_id(self, chapter_num: int, idx: int) -> str:
        """生成全局唯一的chunk_id"""
        # 使用全局计数器确保唯一性
        self.chunk_id_counter += 1
        chunk_id = f"CH{chapter_num}_{self.chunk_id_counter:03d}"
        
        # 检查重复（理论上不应该发生）
        while chunk_id in self.used_chunk_ids:
            self.chunk_id_counter += 1
            chunk_id = f"CH{chapter_num}_{self.chunk_id_counter:03d}"
        
        self.used_chunk_ids.add(chunk_id)
        return chunk_id
    
    def process_file(self, filepath: Path) -> List[Dict]:
        """
        处理单个文件
        
        流程：
        1. 数据清洗：加载并清洗文档
        2. 结构切分：按章节/子标题/条款切分
        3. 语义切分：基于相似度合并/拆分
        4. 生成JSON输出
        """
        print(f"\n{'='*60}")
        print(f"处理文件: {filepath.name}")
        print(f"{'='*60}")
        
        # 步骤1：数据清洗
        print("\n[步骤1] 数据清洗...")
        try:
            text, tables, table_indices = self.cleaner.load_document(filepath)
            print(f"  ✓ 文档加载完成，共 {len(text)} 字符")
            print(f"  ✓ 提取到 {len(tables)} 个表格")
        except Exception as e:
            print(f"  ✗ 文档加载失败: {e}")
            return []
        
        # 步骤2：结构切分
        print("\n[步骤2] 结构切分...")
        units = self.parser.parse_structure(text, tables, table_indices)
        print(f"  ✓ 结构切分完成，共 {len(units)} 个单元")
        
        # 按章节分组
        chapter_groups = {}
        for unit in units:
            chapter = unit.chapter if unit.chapter else "未分类"
            if chapter not in chapter_groups:
                chapter_groups[chapter] = []
            chapter_groups[chapter].append(unit)
        
        print(f"  ✓ 识别到 {len(chapter_groups)} 个章节")
        
        # 步骤3：语义切分
        print("\n[步骤3] 语义切分...")
        all_chunks = []
        
        for chapter_name, chapter_units in chapter_groups.items():
            chapter_num = self.extract_chapter_number(chapter_name)
            print(f"\n  处理章节: {chapter_name} (编号: {chapter_num})")
            
            # 对该章节的单元进行语义切分
            chunks = self.chunker.semantic_chunking(chapter_units, chapter_num)
            
            # 重新分配全局唯一的chunk_id
            for i, chunk in enumerate(chunks):
                chunk.chunk_id = self.generate_unique_chunk_id(chapter_num, i + 1)
            
            all_chunks.extend(chunks)
            print(f"    → 生成 {len(chunks)} 个chunks")
        
        print(f"\n  ✓ 语义切分完成，共 {len(all_chunks)} 个chunks")
        
        # 步骤4：转换为字典列表
        print("\n[步骤4] 生成JSON输出...")
        result = [chunk.to_dict() for chunk in all_chunks]
        
        # 统计信息
        total_chars = sum(c.metadata["char_count"] for c in all_chunks)
        total_tables = sum(1 for c in all_chunks if c.metadata["has_table"])
        avg_similarity = sum(c.metadata["avg_similarity"] for c in all_chunks) / len(all_chunks) if all_chunks else 0
        
        print(f"  ✓ 总字符数: {total_chars}")
        print(f"  ✓ 包含表格的chunks: {total_tables}")
        print(f"  ✓ 平均语义相似度: {avg_similarity:.3f}")
        
        return result
    
    def process_all(self):
        """处理所有文档"""
        print("\n" + "="*60)
        print("RAG文档处理 - 混合切分")
        print("="*60)
        print(f"输入目录: {self.data_dir}")
        print(f"输出目录: {self.output_dir}")
        print(f"相似度阈值: {self.chunker.similarity_threshold}")
        
        # 支持的文件格式
        supported_extensions = ['.txt', '.doc', '.docx', '.pdf']
        
        # 查找所有文档
        all_files = []
        for ext in supported_extensions:
            all_files.extend(self.data_dir.glob(f"*{ext}"))
        
        if not all_files:
            print(f"\n⚠️ 在 {self.data_dir} 中未找到支持的文档")
            return
        
        print(f"\n发现 {len(all_files)} 个待处理文档")
        
        # 处理每个文件
        for filepath in sorted(all_files):
            chunks = self.process_file(filepath)
            
            if chunks:
                # 保存结果
                output_filename = f"{filepath.stem}_chunks.json"
                output_path = self.output_dir / output_filename
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(chunks, f, ensure_ascii=False, indent=2)
                
                print(f"\n  💾 已保存: {output_path}")
        
        print("\n" + "="*60)
        print("✅ 所有文档处理完成！")
        print(f"   全局Chunk ID范围: CH1_001 ~ CH{len(self.used_chunk_ids)}_{self.chunk_id_counter:03d}")
        print("="*60)


# ==================== 主函数 ====================

def main():
    """主函数"""
    import shutil
    
    # 配置路径
    data_dir = "/root/autodl-tmp/rag/test_chunking/data"
    output_dir = "/root/autodl-tmp/rag/test_chunking/output"
    # 使用完整的模型快照路径
    model_path = "/root/autodl-tmp/rag/backend/data/models/BAAI--bge-base-zh-v1.5/models--BAAI--bge-base-zh-v1.5/snapshots/f03589ceff5aac7111bd60cfc7d497ca17ecac65"
    
    # 清空输出目录（确保chunk_id全局唯一）
    output_path = Path(output_dir)
    if output_path.exists():
        for f in output_path.glob("*_chunks.json"):
            f.unlink()
    
    # 创建处理器并执行
    # 提高阈值到0.95，减少合并，获得更细粒度的切分
    processor = RAGDocumentProcessor(
        data_dir=data_dir,
        output_dir=output_dir,
        model_path=model_path,
        similarity_threshold=0.95
    )
    
    processor.process_all()


if __name__ == "__main__":
    main()
